import pytest
import os
import tempfile
from unittest.mock import patch, MagicMock
from app.formatter import MemoFormatter
from docx import Document

def test_memo_formatter_initialization():
    """Test that the memo formatter initializes correctly."""
    # Test with default output folder
    with patch.dict(os.environ, {}, clear=True):  # Clear all env vars
        with patch.dict(os.environ, {"UPLOAD_FOLDER": "./output"}):  # Set to a valid path
            with patch("os.makedirs") as mock_makedirs:  # Mock makedirs to avoid file system operations
                formatter = MemoFormatter()
                assert formatter.output_folder == "./output"
                mock_makedirs.assert_called_once_with("./output", exist_ok=True)

    # Test with custom output folder
    with tempfile.TemporaryDirectory() as temp_dir:
        with patch.dict(os.environ, {"UPLOAD_FOLDER": temp_dir}):
            with patch("os.makedirs") as mock_makedirs:  # Mock makedirs to avoid file system operations
                formatter = MemoFormatter()
                assert formatter.output_folder == temp_dir
                mock_makedirs.assert_called_once_with(temp_dir, exist_ok=True)

    # Test that the output folder is created if it doesn't exist
    with tempfile.TemporaryDirectory() as temp_dir:
        test_path = os.path.join(temp_dir, "test_output")
        with patch.dict(os.environ, {"UPLOAD_FOLDER": test_path}):
            with patch("os.makedirs") as mock_makedirs:
                formatter = MemoFormatter()
                mock_makedirs.assert_called_once_with(test_path, exist_ok=True)

def test_format_memo_as_docx():
    """Test formatting a memo as a Word document."""
    # Create a temporary directory for the test
    with tempfile.TemporaryDirectory() as temp_dir:
        # Set up the formatter with the temporary directory
        with patch.dict(os.environ, {"UPLOAD_FOLDER": temp_dir}):
            formatter = MemoFormatter()

            # Test memo text
            memo_text = """
            Executive Summary

            This is a test memo for Company XYZ.

            1. Financial Highlights

            Revenue: $1,000,000
            EBITDA: $200,000

            2. Risk Analysis

            Low risk profile.
            """

            # Mock datetime to get a consistent filename
            with patch("app.formatter.datetime") as mock_datetime:
                mock_now = MagicMock()
                mock_now.strftime.return_value = "20230101"
                mock_datetime.now.return_value = mock_now

                # Format the memo
                file_path = formatter.format_memo_as_docx(memo_text, "Test Company")

                # Check that the file was created
                assert os.path.exists(file_path)
                assert file_path.endswith(".docx")
                assert "Test_Company" in file_path

                # Check the content of the document
                doc = Document(file_path)

                # Check that the document has the expected number of paragraphs
                # (Title + Date + Company + Separator + Memo paragraphs)
                assert len(doc.paragraphs) > 4

                # Check the title
                assert "Credit Memo" in doc.paragraphs[0].text

                # Check the company name
                assert "Test Company" in doc.paragraphs[2].text

                # Check that the memo text is included
                full_text = "\n".join([p.text for p in doc.paragraphs])
                assert "Executive Summary" in full_text
                assert "Financial Highlights" in full_text
                assert "Risk Analysis" in full_text

def test_format_memo_with_special_characters():
    """Test formatting a memo with special characters in the company name."""
    # Create a temporary directory for the test
    with tempfile.TemporaryDirectory() as temp_dir:
        # Set up the formatter with the temporary directory
        with patch.dict(os.environ, {"UPLOAD_FOLDER": temp_dir}):
            formatter = MemoFormatter()

            # Test with special characters in company name
            special_company_name = "Company/With:Special*Characters?"

            # Mock datetime to get a consistent filename
            with patch("app.formatter.datetime") as mock_datetime:
                mock_now = MagicMock()
                mock_now.strftime.return_value = "20230101"
                mock_datetime.now.return_value = mock_now

                # Format the memo
                file_path = formatter.format_memo_as_docx("Test memo content", special_company_name)

                # Check that the file was created with sanitized name
                assert os.path.exists(file_path)
                assert "Company_With_Special_Characters_" in file_path

                # Check that the original company name is in the document
                doc = Document(file_path)
                company_paragraph = doc.paragraphs[2].text
                assert special_company_name in company_paragraph

def test_document_properties():
    """Test that the document properties are set correctly."""
    # Create a temporary directory for the test
    with tempfile.TemporaryDirectory() as temp_dir:
        # Set up the formatter with the temporary directory
        with patch.dict(os.environ, {"UPLOAD_FOLDER": temp_dir}):
            formatter = MemoFormatter()

            # Create a mock document
            doc = Document()

            # Call the private method to set document properties
            formatter._set_document_properties(doc)

            # Check that the document has sections
            assert len(doc.sections) > 0

            # Check that the normal style is set
            assert 'Normal' in doc.styles
            normal_style = doc.styles['Normal']
            assert normal_style.font.name == 'Calibri'
