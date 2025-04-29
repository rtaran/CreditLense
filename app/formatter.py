from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

class MemoFormatter:
    """
    Formats credit memos as Word documents.
    """
    
    def __init__(self):
        """Initialize the memo formatter."""
        self.output_folder = os.getenv("UPLOAD_FOLDER", "./output")
        os.makedirs(self.output_folder, exist_ok=True)
    
    def format_memo_as_docx(self, memo_text: str, company_name: str = "Company") -> str:
        """
        Format a credit memo as a Word document.
        
        Args:
            memo_text: The text content of the credit memo
            company_name: The name of the company (for the filename)
            
        Returns:
            The path to the generated Word document
        """
        # Create a new Document
        doc = Document()
        
        # Set document properties
        self._set_document_properties(doc)
        
        # Add title
        title = doc.add_heading("Credit Memo", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_paragraph.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(10)
        
        # Add company name
        company_paragraph = doc.add_paragraph()
        company_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        company_run = company_paragraph.add_run(f"Company: {company_name}")
        company_run.font.bold = True
        
        # Add a horizontal line
        doc.add_paragraph("_" * 50)
        
        # Process the memo text
        sections = memo_text.split("\n\n")
        for section in sections:
            if section.strip():
                # Check if this is a heading (starts with number followed by period)
                if section.strip()[0].isdigit() and "." in section.split()[0]:
                    # This is likely a heading
                    doc.add_heading(section.strip(), level=2)
                else:
                    # Regular paragraph
                    doc.add_paragraph(section.strip())
        
        # Save the document
        sanitized_company_name = "".join(c if c.isalnum() else "_" for c in company_name)
        filename = f"credit_memo_{sanitized_company_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        file_path = os.path.join(self.output_folder, filename)
        doc.save(file_path)
        
        return file_path
    
    def _set_document_properties(self, doc):
        """Set document properties like margins, styles, etc."""
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Set styles
        styles = doc.styles
        style = styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

# Singleton instance for easy import
memo_formatter = MemoFormatter()